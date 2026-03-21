# -*- coding: utf-8 -*-
"""
ask_govpersona.py  —  GovPersona CLI for restricted work computers.

Ask any Israeli government agency a question and receive a Word document (.docx).
No browser, no UI, no admin rights needed — just Python 3.8+ and two pip packages.

FIRST-TIME SETUP (run ONCE on the work computer):
  install.bat          installs anthropic + python-docx from local whl_packages/

USAGE:
  python ask_govpersona.py --org finance_ministry --question "מה מדיניות הממשלה בנושא הגירעון?"
  python ask_govpersona.py --org finance_ministry -q "What is Israel's deficit target for 2025?" -o report.docx
  ask.bat --org central_bank -q "What is the current interest rate policy?"

AGENTS:
  finance_ministry      Ministry of Finance (משרד האוצר)
  central_bank          Bank of Israel (בנק ישראל)
  securities_authority  Israel Securities Authority (רשות ניירות ערך)
"""
import sys
import io

# Fix Windows CP1255 encoding so Hebrew text doesn't crash print()
if sys.stdout.encoding and sys.stdout.encoding.upper() not in ("UTF-8", "UTF8"):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8", errors="replace")

import os
import re
import json
import argparse
import datetime

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))

_KNOWN_AGENTS = ["finance_ministry", "central_bank", "securities_authority"]


# ── .env loader ────────────────────────────────────────────────────────────────

def _load_env():
    """Read ANTHROPIC_API_KEY from .env file next to this script."""
    env_file = os.path.join(SCRIPT_DIR, ".env")
    if not os.path.exists(env_file):
        return
    with open(env_file, encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if line and not line.startswith("#") and "=" in line:
                k, v = line.split("=", 1)
                os.environ.setdefault(k.strip(), v.strip())


# ── Keyword retrieval (pure Python, no embeddings) ─────────────────────────────

def _tokenize(text: str) -> set:
    """Lowercase, remove punctuation, return set of tokens with len >= 2."""
    return {t for t in re.sub(r"[^\w\s]", " ", text.lower()).split() if len(t) >= 2}


def keyword_search(chunks: list, query: str, top_k: int = 7) -> list:
    """
    Score chunks by how many query tokens appear in them.
    Falls back to first top_k chunks if nothing matches.
    """
    query_tokens = _tokenize(query)
    if not query_tokens:
        return chunks[:top_k]

    scored = []
    for chunk in chunks:
        text_lower = chunk["text"].lower()
        score = sum(1 for t in query_tokens if t in text_lower)
        scored.append((score, chunk))

    scored.sort(key=lambda x: -x[0])
    return [c for _, c in scored[:top_k]]


# ── System prompt (exact copy of persona_engine._SYSTEM_TEMPLATE) ─────────────

_SYSTEM_TEMPLATE = """\
You are a senior official and institutional spokesperson for {org_name}, with deep expertise \
in its mandate, analytical frameworks, research tradition, and known policy positions.

YOUR TWO SOURCES OF KNOWLEDGE — use both, in order of preference:

1. UPLOADED DOCUMENTS (retrieved context below): When the answer is found here, cite the source \
and present it as the organization's documented position.

2. INSTITUTIONAL KNOWLEDGE: When uploaded documents do not cover the topic, draw on your \
broad knowledge of {org_name}'s known positions, economic principles it upholds, its past \
research and publications, and the analytical frameworks it applies. Senior officials are \
expected to articulate well-reasoned positions — not simply say "no document covers this."

HOW TO HANDLE EACH CASE:
- If the retrieved context directly answers the question → cite it and answer from it.
- If the retrieved context is partially relevant → use it as a starting point, then extend \
with institutional reasoning. Note which parts are from documents vs. your analysis.
- If the retrieved context is not relevant → answer entirely from institutional knowledge. \
Open with a brief phrase such as "Based on {org_name}'s known analytical framework..." or \
"From {org_name}'s perspective..." to signal you are reasoning from institutional knowledge \
rather than a specific uploaded document.
- NEVER refuse to engage with a substantive policy question. A senior {role_title} always \
has a view — even if it is nuanced or conditional.
- Do NOT fabricate specific statistics, interest-rate decisions, or named publications \
that you cannot verify. Reason qualitatively when precise data is unavailable.

TONE AND LANGUAGE:
- Adopt the formal, professional register of a senior {role_title} from {org_name}.
- Respond in the SAME LANGUAGE as the user's question (Hebrew \u2192 Hebrew, English \u2192 English).
- Use {org_name}'s official terminology and analytical vocabulary.

MANDATE: {org_mandate}

RETRIEVED CONTEXT FROM UPLOADED DOCUMENTS:
{context}
"""


def _build_context(top_chunks: list) -> str:
    if not top_chunks:
        return "[No documents found for this agency.]"
    parts = []
    for i, chunk in enumerate(top_chunks, 1):
        source = chunk.get("source", "unknown")
        parts.append(f"[Source {i}: {source}]\n{chunk['text']}")
    return "\n\n---\n\n".join(parts)


def _build_system_prompt(cfg: dict, context: str) -> str:
    return _SYSTEM_TEMPLATE.format(
        org_name=cfg["name"],
        role_title=cfg["role_title"],
        org_mandate=cfg["org_mandate"],
        context=context,
    )


# ── Word document builder ──────────────────────────────────────────────────────

def _save_docx(question: str, answer: str, sources: list, cfg: dict, out_path: str):
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
    except ImportError:
        print("\nERROR: python-docx is not installed.")
        print("  Run install.bat once to install it from the whl_packages/ folder.")
        sys.exit(1)

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)

    # ── Title block ───────────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(cfg.get("name_en", cfg["name"]))
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    role_para = doc.add_paragraph()
    role_run = role_para.add_run(cfg.get("role_title", ""))
    role_run.italic = True
    role_run.font.size = Pt(11)
    role_run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    date_para = doc.add_paragraph()
    date_run = date_para.add_run(datetime.date.today().strftime("%B %d, %Y"))
    date_run.font.size = Pt(10)
    date_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph()

    # ── Question ──────────────────────────────────────────────────────────────
    qh = doc.add_heading("Question", level=2)
    if qh.runs:
        qh.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    q_para = doc.add_paragraph(question)
    q_para.runs[0].italic = True
    q_para.runs[0].font.size = Pt(11)

    doc.add_paragraph()

    # ── Answer ────────────────────────────────────────────────────────────────
    ah = doc.add_heading("Answer", level=2)
    if ah.runs:
        ah.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    for block in answer.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        lines = block.split("\n")
        # Detect bullet-like blocks (all non-empty lines start with - * or bullet)
        bullet_lines = [l.strip() for l in lines if l.strip()]
        is_bullet_block = bullet_lines and all(
            l.startswith(("-", "*", "\u2022", "\u25e6")) for l in bullet_lines
        )
        if is_bullet_block:
            for line in bullet_lines:
                line_text = line.lstrip("-*\u2022\u25e6 ").strip()
                if line_text:
                    p = doc.add_paragraph(line_text, style="List Bullet")
                    if p.runs:
                        p.runs[0].font.size = Pt(11)
        else:
            p = doc.add_paragraph(block)
            if p.runs:
                p.runs[0].font.size = Pt(11)

    # ── Sources ───────────────────────────────────────────────────────────────
    unique_sources = sorted(set(sources))
    if unique_sources:
        doc.add_paragraph()
        sh = doc.add_heading("Sources Consulted", level=2)
        if sh.runs:
            sh.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        for src in unique_sources:
            p = doc.add_paragraph(src, style="List Bullet")
            if p.runs:
                p.runs[0].font.size = Pt(10)
                p.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    # ── Footer note ───────────────────────────────────────────────────────────
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_run = footer_para.add_run(
        "Generated by GovPersona  |  Powered by Claude (Anthropic)"
    )
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    footer_run.italic = True

    doc.save(out_path)
    print(f"\n  Saved: {out_path}")


# ── Main ───────────────────────────────────────────────────────────────────────

def main():
    _load_env()

    parser = argparse.ArgumentParser(
        description="Ask a GovPersona agent a question and get a Word document answer.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=__doc__,
    )
    parser.add_argument(
        "--org",
        required=True,
        choices=_KNOWN_AGENTS,
        metavar="AGENT",
        help="Agent to ask: " + ", ".join(_KNOWN_AGENTS),
    )
    parser.add_argument(
        "--question", "-q",
        required=True,
        help="Your question in Hebrew or English",
    )
    parser.add_argument(
        "--out", "-o",
        default=None,
        help="Output .docx path (default: answer_<org>_<date>.docx)",
    )
    parser.add_argument(
        "--top-k",
        type=int,
        default=7,
        help="Number of document chunks to retrieve (default: 7)",
    )
    args = parser.parse_args()

    # Default output filename
    if not args.out:
        date_str = datetime.date.today().strftime("%Y-%m-%d")
        args.out = os.path.join(SCRIPT_DIR, f"answer_{args.org}_{date_str}.docx")

    print()
    print("  GovPersona CLI")
    print(f"  Agent   : {args.org}")
    print(f"  Question: {args.question[:80]}{'...' if len(args.question) > 80 else ''}")
    print(f"  Output  : {args.out}")
    print()

    # ── Load agent config ─────────────────────────────────────────────────────
    agents_file = os.path.join(SCRIPT_DIR, "agents_config.json")
    if not os.path.exists(agents_file):
        print(f"ERROR: agents_config.json not found at:\n  {agents_file}")
        print("  Run export_kb.py on your home laptop first.")
        sys.exit(1)

    with open(agents_file, encoding="utf-8") as f:
        agents = json.load(f)

    if args.org not in agents:
        print(f"ERROR: Agent '{args.org}' not found in agents_config.json")
        print(f"  Available agents: {', '.join(agents.keys())}")
        sys.exit(1)

    cfg = agents[args.org]

    # ── Load knowledge base ───────────────────────────────────────────────────
    kb_file = os.path.join(SCRIPT_DIR, f"kb_{args.org}.json")
    if not os.path.exists(kb_file):
        print(f"ERROR: Knowledge base file not found:\n  {kb_file}")
        print("  Run export_kb.py on your home laptop first, then copy the JSON file here.")
        sys.exit(1)

    print("  Loading knowledge base...", end=" ", flush=True)
    with open(kb_file, encoding="utf-8") as f:
        chunks = json.load(f)
    print(f"{len(chunks):,} chunks loaded")

    # ── Retrieve relevant chunks ──────────────────────────────────────────────
    print("  Searching for relevant context...", end=" ", flush=True)
    top_chunks = keyword_search(chunks, args.question, top_k=args.top_k)
    sources = [c.get("source", "unknown") for c in top_chunks]
    print(f"{len(top_chunks)} chunks selected")
    for c in top_chunks:
        print(f"    - {c.get('source', '?')}  (chunk {c.get('chunk_index', '?')})")

    # ── Build prompt ──────────────────────────────────────────────────────────
    context = _build_context(top_chunks)
    system_prompt = _build_system_prompt(cfg, context)

    # ── Check API key ─────────────────────────────────────────────────────────
    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        print("\nERROR: ANTHROPIC_API_KEY is not set.")
        print("  Add the following line to the .env file next to this script:")
        print("  ANTHROPIC_API_KEY=sk-ant-...")
        sys.exit(1)

    # ── Call Claude ───────────────────────────────────────────────────────────
    try:
        import anthropic
    except ImportError:
        print("\nERROR: The 'anthropic' package is not installed.")
        print("  Run install.bat to install it from the whl_packages/ folder.")
        sys.exit(1)

    print("\n  Calling Claude API...", end=" ", flush=True)
    client = anthropic.Anthropic(api_key=api_key)
    response = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=2000,
        temperature=0.3,
        system=system_prompt,
        messages=[{"role": "user", "content": args.question}],
    )
    answer = response.content[0].text
    print("done")

    # ── Save Word document ────────────────────────────────────────────────────
    _save_docx(args.question, answer, sources, cfg, args.out)

    print(f"\n  Open  {os.path.basename(args.out)}  in Word to read the answer.")
    print("  (Double-click the file in File Explorer)\n")


if __name__ == "__main__":
    main()
